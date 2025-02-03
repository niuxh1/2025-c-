
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

def create_glossary_word_doc():
    """创建包含符号意义表的 Word 文档."""

    document = Document()

    # 添加 "Glossary & Symbols" 标题 (二级标题)
    document.add_heading('4. Glossary & Symbols', level=2)

    # 表格数据 (从之前的 HTML 表格中提取)
    table_data = [
        ("Symbol", "Meaning"),
        (r"\( y \)", "目标变量 (如奖牌数)"),
        (r"\( \hat{y} \)", "模型的预测值 (随机森林, XGBoost)"),
        (r"\( \hat{y}_{\text{stack}} \)", "Stacking 模型的最终预测值"),
        (r"\( \hat{y}_{\text{Poisson}} \)", "Poisson 回归模型的预测值"),
        (r"\( \hat{y}_{\text{RF}} \)", "随机森林模型的预测值"),
        (r"\( X \)", r"特征向量，定义为 \( (x_{\text{noc},1}, x_{\text{noc},2}, \dots, x_{\text{noc},k}, x_{\text{host}}, x_{\text{year}}, x_{\text{athletes}}, x_{\text{events}}) \)"),
        (r"\( x_{\text{noc},j} \)", r"第 \( j \) 个 NOC（国家奥委会）类别的热编码变量（\( j = 1, 2, \dots, k \)）"),
        (r"\( x_{\text{host}} \)", "二元变量，表示是否为主办国（1 表示是主办国，0 表示不是）"),
        (r"\( x_{\text{year}} \)", "年份变量，表示奥运会举办的年份"),
        (r"\( x_{\text{athletes}} \)", "运动员总数变量，表示某国家在某届奥运会中参赛的运动员总数"),
        (r"\( x_{\text{events}} \)", "参加项目总数变量，表示某国家在某届奥运会中参加的项目总数"),
        (r"\( k \)", "NOC 类别的总数"),
        (r"\( \beta_0 \)", "Poisson 回归模型的截距项"),
        (r"\( \beta_{\text{noc},j} \)", r"第 \( j \) 个 NOC 类别的系数 (Poisson 回归模型)"),
        (r"\( \beta_{\text{host}} \)", "<code>host</code> 变量的系数 (Poisson 回归模型)"),
        (r"\( \beta_{\text{year}} \)", "<code>year</code> 变量的系数 (Poisson 回归模型)"),
        (r"\( \beta_{\text{athletes}} \)", "<code>athletes</code> 变量的系数 (Poisson 回归模型)"),
        (r"\( \beta_{\text{events}} \)", "<code>events</code> 变量的系数 (Poisson 回归模型)"),
        (r"\( T \)", "决策树的总数 (随机森林, XGBoost)"),
        (r"\( f_t(X) \)", "第 \( t \) 棵树的预测函数 (随机森林, XGBoost)"),
        (r"\( g(\cdot) \)", "元模型 (Stacking 模型)，用于结合基模型的预测结果"),
        (r"\( L \)", "XGBoost 模型的 目标函数"),
        (r"\( l(y_i, \hat{y}_i) \)", r"损失函数，用于衡量真实值 \( y_i \) 和预测值 \( \hat{y}_i \) 之间的差异 (XGBoost 模型)"),
        (r"\( \Omega(f_t) \)", r"正则化项，用于控制第 \( t \) 棵树的复杂度 (XGBoost 模型)"),
        (r"\( E[y \mid X] \)", r"给定特征向量 \(X\) 时，目标变量 \(y\) 的期望值 (Poisson 回归模型)")
    ]

    # 添加表格
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # 使用网格表格样式

    # 填充表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "符号 (Symbol)"
    hdr_cells[1].text = "意义 (Meaning)"

    # 填充表格内容
    for symbol, meaning in table_data[1:]: # 跳过表头行
        row_cells = table.add_row().cells
        row_cells[0].text = symbol
        row_cells[1].text = meaning

    # 保存 Word 文档
    document.save('Glossary_and_Symbols.docx')
    print("Word 文档 'Glossary_and_Symbols.docx' 创建成功！")

if __name__ == "__main__":
    create_glossary_word_doc()
